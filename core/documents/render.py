"""Render resume / cover-letter text into PDF, DOCX, Markdown, and HTML.

Input is lightweight Markdown-ish plain text (lines, ``#``/``##`` headings,
``-``/``*`` bullets, blank lines as spacers). Each renderer writes a file and
returns its path.
"""

from __future__ import annotations

from pathlib import Path
from xml.sax.saxutils import escape

from core.database.enums import DocumentFormat

_EXTENSIONS = {
    DocumentFormat.pdf: "pdf",
    DocumentFormat.docx: "docx",
    DocumentFormat.markdown: "md",
    DocumentFormat.html: "html",
}


def extension_for(fmt: DocumentFormat) -> str:
    return _EXTENSIONS[fmt]


def render_html(title: str, content: str) -> str:
    parts = []
    for raw in content.splitlines():
        line = raw.rstrip()
        if not line.strip():
            parts.append("")
            continue
        if line.startswith("## "):
            parts.append(f"<h3>{escape(line[3:].strip())}</h3>")
        elif line.startswith("# "):
            parts.append(f"<h2>{escape(line[2:].strip())}</h2>")
        elif line.lstrip().startswith(("- ", "* ")):
            parts.append(f"<li>{escape(line.lstrip()[2:].strip())}</li>")
        else:
            parts.append(f"<p>{escape(line)}</p>")
    body = "\n".join(parts)
    return (
        "<!doctype html>\n<html lang='en'>\n<head>\n"
        "<meta charset='utf-8'>\n"
        f"<title>{escape(title)}</title>\n"
        "<style>body{font-family:system-ui,Arial,sans-serif;max-width:48rem;"
        "margin:2rem auto;line-height:1.5;padding:0 1rem}h1{margin-bottom:0}"
        "li{margin-left:1rem}</style>\n</head>\n<body>\n"
        f"<h1>{escape(title)}</h1>\n{body}\n</body>\n</html>\n"
    )


def _write_markdown(path: Path, title: str, content: str) -> None:
    path.write_text(f"# {title}\n\n{content}\n", encoding="utf-8")


def _write_html(path: Path, title: str, content: str) -> None:
    path.write_text(render_html(title, content), encoding="utf-8")


def _write_pdf(path: Path, title: str, content: str) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    flow: list = [Paragraph(escape(title), styles["Title"]), Spacer(1, 12)]
    for raw in content.splitlines():
        line = raw.strip()
        if not line:
            flow.append(Spacer(1, 6))
            continue
        if line.startswith("## "):
            flow.append(Paragraph(escape(line[3:].strip()), styles["Heading3"]))
        elif line.startswith("# "):
            flow.append(Paragraph(escape(line[2:].strip()), styles["Heading2"]))
        elif line.startswith(("- ", "* ")):
            flow.append(Paragraph(f"• {escape(line[2:].strip())}", styles["BodyText"]))
        else:
            flow.append(Paragraph(escape(line), styles["BodyText"]))
    SimpleDocTemplate(str(path), pagesize=letter).build(flow)


def _write_docx(path: Path, title: str, content: str) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    for raw in content.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=3)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=2)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(str(path))


def export(content: str, *, fmt: DocumentFormat, path: Path, title: str) -> Path:
    """Render ``content`` to ``path`` in the given format. Returns the path."""
    path.parent.mkdir(parents=True, exist_ok=True)
    writer = {
        DocumentFormat.markdown: _write_markdown,
        DocumentFormat.html: _write_html,
        DocumentFormat.pdf: _write_pdf,
        DocumentFormat.docx: _write_docx,
    }[fmt]
    writer(path, title, content)
    return path
